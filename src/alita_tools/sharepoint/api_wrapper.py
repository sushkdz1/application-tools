import logging
import os
from io import BytesIO
from traceback import format_exc
import json
from typing import List, Optional, Any, Dict
from langchain_core.tools import ToolException
from langchain_core.pydantic_v1 import root_validator, BaseModel
from pydantic import create_model
from pydantic.fields import FieldInfo
from docx import Document

logger = logging.getLogger(__name__)

NoInput = create_model(
    "NoInput"
)

SharepointReadList = create_model(
    "SharepointSearchModel",
    list_title=(str, FieldInfo(description="Name of a Sharepoint list to be read."))
)

SharepointGetFilesInFolder = create_model(
    "SharepointGetFilesInFolder",
    folder=(str, FieldInfo(description="Folder name to get list of the files."))
)

SharepointReadDocument = create_model(
    "SharepointReadDocument",
    path=(str, FieldInfo(description="Contains the server-relative path of the  for reading."))
)


def read_docx(path):
    """Reads and prints content from a .docx file."""
    try:
        doc = Document(path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except Exception as e:
        print(f"Error reading {path}: {e}")
        return ""

def read_docx_from_bytes(file_content):
    """Read and return content from a .docx file using a byte stream."""
    try:
        doc = Document(BytesIO(file_content))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except Exception as e:
        print(f"Error reading .docx from bytes: {e}")
        return ""

class SharepointApiWrapper(BaseModel):
    site_url: str
    list_title: Optional[str]
    client_id: str
    client_secret: str

    @root_validator()
    def validate_toolkit(cls, values):

        try:
            from office365.runtime.auth.authentication_context import AuthenticationContext
            from office365.sharepoint.client_context import ClientContext
        except ImportError:
            raise ImportError(
                "`office365` package not found, please run "
               "`pip install office365` and `pip install office365-rest-python-client`"
            )

        site_url = values['site_url']
        #list_title = values.get('list_title')
        client_id = values.get('client_id')
        client_secret = values.get('client_secret')

        values['client'] = None
        try:
            ctx_auth = AuthenticationContext(site_url)
            if ctx_auth.acquire_token_for_app(client_id, client_secret):
                values['client'] = ClientContext(site_url, ctx_auth)
            else:
                logging.error("Failed to authenticate with SharePoint.")
        except Exception as e:
                logging.error(f"Failed to authenticate with SharePoint: {str(e)}")

        return values


    def read_list(self, list_title):
        """ Reads a specified List in sharepoint site """
        try:
            target_list = self.client.web.lists.get_by_title(list_title)
            self.client.load(target_list)
            self.client.execute_query()
            items = target_list.items.get().top(1000).execute_query()
            logging.info("{0} items from sharepoint loaded successfully.".format(len(items)))
            result = []
            for item in items:
                result.append(item.properties)
            return result
        except Exception as e:
            logging.error(f"Failed to load items from sharepoint: {e}")

    def get_all_files(self):
        """ Lists all files from sharepoint."""
        try:
            doc_lib = self.client.web.lists.get_by_title('Documents')
            items = doc_lib.items
            self.client.load(items).execute_query()
            result = []
            for item in items:
                if item.file_system_object_type == 0: #FileSystemObjectType.File
                    file = item.file
                    self.client.load(file).execute_query()
                    temp_props = {'Name': file.properties['Name'],
                                 'Path': file.properties['ServerRelativeUrl'],
                                 'Created': file.properties['TimeCreated'],
                                 'Modified' : file.properties['TimeLastModified'],
                                 'Link': file.properties['LinkingUrl']
                                 }
                    result.append(temp_props)
            #print(result)
            return result
        except Exception as e:
            logging.error(f"Failed to load files from sharepoint: {e}")

    def get_all_files_in_folder(self, folder):
        """ Lists all files from sharepoint in a specific folder."""
        try:
            target_folder_url = "Shared Documents/" + folder
            root_folder = self.client.web.get_folder_by_server_relative_path(target_folder_url)
            files = root_folder.get_files(True).execute_query()
            result = []
            for file in files:
                temp_props = {'Name': file.properties['Name'],
                              'Path': file.properties['ServerRelativeUrl'],
                              'Created': file.properties['TimeCreated'],
                              'Modified': file.properties['TimeLastModified'],
                              'Link': file.properties['LinkingUrl']
                              }
                result.append(temp_props)
            #print(result)
            return result
        except Exception as e:
            logging.error(f"Failed to load files from sharepoint: {e}")

    def read_file(self, path):
        """ Reads file located at the specified server-relative path """
        file = self.client.web.get_file_by_server_relative_path(path)
        self.client.load(file)
        self.client.execute_query()

        file_content = file.read()
        self.client.execute_query()

        if file.name.endswith('.txt'):
            try:
                file_content_str = file_content.decode('utf-8')
                print(file_content_str)
            except Exception as e:
                print(f"Error decoding file content: {e}")
        elif file.name.endswith('.docx'):
            file_content_str = read_docx_from_bytes(file_content)
        else:
            return "Not supported type of files entered. Supported types are TXT and DOCX only at the moment"
        return file_content_str

    def get_available_tools(self):
        return [
            {
                "name": "read_list",
                "description": self.read_list.__doc__,
                "args_schema": SharepointReadList,
                "ref": self.read_list
            },
            {
                "name": "get_all_files",
                "description": self.get_all_files.__doc__,
                "args_schema": NoInput,
                "ref": self.get_all_files
            },
            {
                "name": "get_all_files_in_folder",
                "description": self.get_all_files_in_folder.__doc__,
                "args_schema": SharepointGetFilesInFolder,
                "ref": self.get_all_files_in_folder
            },
            {
                "name": "read_document",
                "description": self.read_file.__doc__,
                "args_schema": SharepointReadDocument,
                "ref": self.read_file
            }
        ]

    def run(self, mode: str, *args: Any, **kwargs: Any):
        for tool in self.get_available_tools():
            if tool["name"] == mode:
                return tool["ref"](*args, **kwargs)
        else:
            raise ValueError(f"Unknown mode: {mode}")
